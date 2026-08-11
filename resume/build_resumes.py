from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
DIST = ROOT / "dist"
DIST.mkdir(parents=True, exist_ok=True)

NAME = "Touhidul Alam Seyam"
CONTACT = (
    "Chattogram, Bangladesh  |  +880 1311-104804  |  seyamalam41@gmail.com  |  "
    "seyamalam.vercel.app  |  github.com/Seyamalam  |  linkedin.com/in/touhidulalamseyam  |  "
    "orcid.org/0009-0007-7512-1893"
)

LEADERSHIP = "Assistant General Secretary | BGCTUB IT Club | Present - Help lead student technology activities, community programs, and club operations."
PRIMARY_WINS = [
    "Champion, HackFusion InnoNation National Hackathon (solo Team Huntrix), 2026",
    "Champion, MicroOps Hackathon at CUET CSE IT Fest (with Abtahee Kabir), 2025",
    "Winner, RoboFusion 1.0 Techathon, 2026",
]
ADDITIONAL_HONORS = [
    "5th Place, 12th IUT ICT Fest GameJam, 2026; 5th Place, IIUC NextGen Hackathon, 2025",
    "Finalist, The Infinity AI BuildFest; Best Campus Ambassador, SciBlitz 2.0",
]

PROJECT_SETS = {
    "software": [
        ("Paris Summit Platform", "Next.js, TypeScript, Convex, Better Auth", "Designed ten responsive directions, then delivered the selected identity as a 19-route client platform with editorial tooling, protected operations, files, forms, and launch-safe donation flows.", "https://github.com/Seyamalam/paris-hindu-summit-2026"),
        ("Huntrix Delta", "Expo, Go, gRPC, Offline-first", "Built the solo HackFusion champion with CRDT inventory sync, signed proof of delivery, offline maps, device-to-device handoff, routing, and reproducible chaos testing.", "https://github.com/Seyamalam/hackfusion_huntrix"),
        ("MicroOps", "Node.js, Docker, S3, CI/CD", "Co-built the CUET champion: a non-blocking file-delivery system with durable jobs, S3-compatible storage, containerized infrastructure, status streaming, and automated delivery checks.", "https://github.com/Seyamalam/cuet-hackathon-1"),
        ("bun-scikit", "TypeScript, Bun, Zig", "Built a scikit-learn-inspired ML library with native Zig acceleration, 209 tracked runtime exports, benchmark gates, and broad preprocessing, model-selection, ensemble, clustering, and metrics APIs.", "https://github.com/Seyamalam/bun-scikit"),
    ],
    "research": [
        ("bun-scikit", "TypeScript, Bun, Zig", "Implemented native-accelerated classical ML with parity contracts, reproducible benchmarks, documentation coverage checks, and 209 tracked runtime exports.", "https://github.com/Seyamalam/bun-scikit"),
        ("Kaggriculture Autonomous Agent", "Python, Simulation, Evaluation", "Engineered a deterministic farming agent and evaluation harness with seeded seat-swapped tournaments, frozen replay corpora, manifest hashing, loss attribution, and regression gates.", "https://github.com/Seyamalam/Kaggriculture"),
        ("Fold-safe ML Pipelines", "Python, LightGBM, XGBoost, CatBoost", "Built leakage-safe five-fold out-of-fold pipelines with nested target encoding, matched-fold screening, schema checks, rank blending, and a 0.96919 public ROC AUC submission.", "https://github.com/Seyamalam/playground-series-s6e8"),
        ("Huntrix Delta", "Expo, Go, gRPC, Offline AI", "Built autonomous triage and predictive route-decay features inside an offline-first disaster logistics system, supported by model-card methodology and a reproducible chaos simulator.", "https://github.com/Seyamalam/hackfusion_huntrix"),
    ],
    "fullstack": [
        ("Paris Summit Platform", "Next.js, Convex, Better Auth", "Turned ten live design directions into a 19-route client platform with structured editing, protected team access, files, forms, and production delivery workflows.", "https://github.com/Seyamalam/paris-hindu-summit-2026"),
        ("ASRRO Portal", "Next.js, TypeScript, Convex", "Built public pages and role-aware operations for membership approvals, events, attendance, content, reports, notifications, and restricted finance access.", "https://github.com/Seyamalam/asrro"),
        ("CareerPath", "Next.js 16, PostgreSQL, AI SDK", "Built profiles, job discovery, explainable match scoring, skill-gap analysis, course recommendations, CV assistance, and generated career roadmaps.", "https://careerpath-vert.vercel.app"),
        ("MicroOps", "Node.js, Docker, S3, CI/CD", "Co-built the CUET champion around durable background jobs, self-hosted S3-compatible storage, responsive status delivery, and containerized infrastructure.", "https://github.com/Seyamalam/cuet-hackathon-1"),
    ],
    "universal": [
        ("Huntrix Delta", "Expo, Go, gRPC, Offline-first", "Built the solo HackFusion champion across offline identity, CRDT inventory, mesh handoff, signed delivery proof, routing, and chaos testing.", "https://github.com/Seyamalam/hackfusion_huntrix"),
        ("Paris Summit Platform", "Next.js, TypeScript, Convex", "Designed ten responsive concepts and shipped the selected direction as a 19-route client platform with protected editorial operations.", "https://github.com/Seyamalam/paris-hindu-summit-2026"),
        ("bun-scikit", "TypeScript, Bun, Zig", "Built a native-accelerated ML library with 209 tracked runtime exports, benchmark gates, and broad estimator and evaluation APIs.", "https://github.com/Seyamalam/bun-scikit"),
        ("RoboFusion", "Python, FastAPI, WebSockets, ESP32", "Built the winning smart-campus response grid with authenticated ingestion, deterministic risk fusion, persistent incident lifecycles, live dashboards, and load tooling.", "https://github.com/Seyamalam/Robofusion"),
    ],
}


@dataclass(frozen=True)
class Theme:
    key: str
    preset: str
    font: str
    ink: str
    accent: str
    muted: str
    title_size: float
    body_size: float
    margins: float
    rule: bool
    compact: bool = False


CLASSIC = Theme(
    key="Software_Engineer_ATS",
    preset="editorial_engineering",
    font="Arial",
    ink="202020",
    accent="A73526",
    muted="5C5854",
    title_size=28,
    body_size=9.85,
    margins=0.66,
    rule=True,
)

MODERN = Theme(
    key="ML_Research_ATS",
    preset="research_blueprint",
    font="Arial",
    ink="17202A",
    accent="1F4E79",
    muted="536273",
    title_size=27,
    body_size=9.7,
    margins=0.66,
    rule=True,
)

COMPACT = Theme(
    key="Full_Stack_Compact_ATS",
    preset="compact_product_sheet",
    font="Calibri",
    ink="1F2933",
    accent="0D7168",
    muted="5B6770",
    title_size=25,
    body_size=9.15,
    margins=0.55,
    rule=False,
    compact=True,
)

UNIVERSAL = Theme(
    key="Universal_ATS",
    preset="universal_navy",
    font="Arial",
    ink="17202A",
    accent="183B56",
    muted="536273",
    title_size=28,
    body_size=9.55,
    margins=0.62,
    rule=True,
)


def set_cellless_border(paragraph, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    p_pr.append(shading)


def set_run_font(run, theme: Theme, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = theme.font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), theme.font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), theme.font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, theme: Theme) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), theme.accent)
    run_props.append(color)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), theme.font)
    fonts.set(qn("w:hAnsi"), theme.font)
    run_props.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(theme.body_size * 2)))
    run_props.append(size)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_doc(theme: Theme) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(theme.margins)
    section.bottom_margin = Inches(theme.margins)
    section.left_margin = Inches(theme.margins)
    section.right_margin = Inches(theme.margins)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = theme.font
    normal._element.rPr.rFonts.set(qn("w:ascii"), theme.font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), theme.font)
    normal.font.size = Pt(theme.body_size)
    normal.font.color.rgb = RGBColor.from_string(theme.ink)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3 if theme.compact else 4)
    normal.paragraph_format.line_spacing = 1.0 if theme.compact else 1.05

    for name, size in (("Heading 1", 11.5 if theme.compact else 12.5), ("Heading 2", 10.5 if theme.compact else 11.5)):
        style = doc.styles[name]
        style.font.name = theme.font
        style._element.rPr.rFonts.set(qn("w:ascii"), theme.font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), theme.font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(theme.accent)
        style.paragraph_format.space_before = Pt(6 if theme.compact else 8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = theme.font
    bullet._element.rPr.rFonts.set(qn("w:ascii"), theme.font)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), theme.font)
    bullet.font.size = Pt(theme.body_size)
    bullet.paragraph_format.left_indent = Inches(0.19 if theme.compact else 0.24)
    bullet.paragraph_format.first_line_indent = Inches(-0.14 if theme.compact else -0.17)
    bullet.paragraph_format.space_after = Pt(1.5 if theme.compact else 2.5)
    bullet.paragraph_format.line_spacing = 1.0 if theme.compact else 1.05

    if "Role Line" not in [s.name for s in doc.styles]:
        role_style = doc.styles.add_style("Role Line", WD_STYLE_TYPE.PARAGRAPH)
    else:
        role_style = doc.styles["Role Line"]
    role_style.font.name = theme.font
    role_style._element.rPr.rFonts.set(qn("w:ascii"), theme.font)
    role_style._element.rPr.rFonts.set(qn("w:hAnsi"), theme.font)
    role_style.font.size = Pt(theme.body_size + 0.3)
    role_style.paragraph_format.space_before = Pt(3)
    role_style.paragraph_format.space_after = Pt(1)
    role_style.paragraph_format.keep_with_next = True

    return doc


def add_header(doc: Document, theme: Theme, target: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if theme is MODERN else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(NAME)
    set_run_font(run, theme, theme.title_size, theme.accent, True)

    role = doc.add_paragraph()
    role.alignment = p.alignment
    role.paragraph_format.space_after = Pt(4)
    r = role.add_run(target.upper() if theme is MODERN else target)
    set_run_font(r, theme, 9.5 if theme.compact else 10.5, theme.muted, True)

    contact = doc.add_paragraph()
    contact.alignment = p.alignment
    contact.paragraph_format.space_after = Pt(7 if theme.compact else 9)
    contact.paragraph_format.line_spacing = 1.0
    cr = contact.add_run(CONTACT)
    set_run_font(cr, theme, 8.4 if theme.compact else 9.0, theme.muted)
    set_cellless_border(contact, theme.accent, 14 if theme in (CLASSIC, UNIVERSAL) else 8)


def add_section(doc: Document, theme: Theme, title: str) -> None:
    p = doc.add_paragraph(title.upper(), style="Heading 1")
    if theme is CLASSIC:
        set_cellless_border(p, theme.accent, 10)
    elif theme is MODERN:
        set_paragraph_shading(p, "E8F0FA")
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
    elif theme is COMPACT:
        set_paragraph_shading(p, theme.accent)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        for run in p.runs:
            set_run_font(run, theme, color="FFFFFF", bold=True)
    else:
        set_paragraph_shading(p, "EAF0F5")
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)


def add_text(doc: Document, theme: Theme, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, theme, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, theme)
    else:
        r = p.add_run(text)
        set_run_font(r, theme)


def add_role(doc: Document, theme: Theme, role: str, org: str, place: str, dates: str) -> None:
    p = doc.add_paragraph(style="Role Line")
    left = p.add_run(f"{role} | {org}")
    set_run_font(left, theme, bold=True)
    right = p.add_run(f"  |  {place}  |  {dates}")
    set_run_font(right, theme, color=theme.muted)


def add_bullets(doc: Document, theme: Theme, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        r = p.add_run(item)
        set_run_font(r, theme)


def add_project(doc: Document, theme: Theme, name: str, stack: str, text: str, url: str | None = None) -> None:
    p = doc.add_paragraph(style="Role Line")
    r = p.add_run(name)
    set_run_font(r, theme, bold=True)
    s = p.add_run(f" | {stack}")
    set_run_font(s, theme, color=theme.muted)
    if url:
        p.add_run(" | ")
        add_hyperlink(p, url.replace("https://", ""), url, theme)
    add_bullets(doc, theme, [text])


def add_project_set(doc: Document, theme: Theme, key: str) -> None:
    for project in PROJECT_SETS[key]:
        add_project(doc, theme, *project)


def add_publication(doc: Document, theme: Theme, title: str, venue: str, year: str, doi: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    r = p.add_run(f'“{title}.” ')
    set_run_font(r, theme, italic=True)
    is_url = doi.startswith("http")
    v = p.add_run(f"{venue}, {year}. {'Link' if is_url else 'DOI'}: ")
    set_run_font(v, theme)
    add_hyperlink(p, doi.replace("https://", ""), doi if is_url else f"https://doi.org/{doi}", theme)


def add_competition_highlights(doc: Document, theme: Theme, *, compact: bool = False) -> None:
    add_section(doc, theme, "Selected Wins")
    add_bullets(doc, theme, PRIMARY_WINS[:2] if compact else PRIMARY_WINS)


def add_leadership_and_honors(doc: Document, theme: Theme) -> None:
    add_section(doc, theme, "Leadership and Additional Honors")
    add_text(doc, theme, LEADERSHIP)
    add_bullets(doc, theme, ADDITIONAL_HONORS)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def software_resume(theme: Theme) -> Document:
    doc = configure_doc(theme)
    add_header(doc, theme, "Software Engineer | Web, App and Backend Systems")
    add_section(doc, theme, "Professional Summary")
    add_text(doc, theme, "Software engineer building production web applications, backend services, mobile systems, and AI-enabled products across TypeScript, Next.js, Node.js, Python, Go, PostgreSQL, and cloud infrastructure. Ships client platforms, resilient distributed prototypes, and open-source developer tooling from architecture through deployment.")
    add_competition_highlights(doc, theme)

    add_section(doc, theme, "Technical Skills")
    add_text(doc, theme, "Languages: TypeScript, JavaScript, Python, SQL, Go, Zig")
    add_text(doc, theme, "Application Development: React, Next.js, Node.js, Django, FastAPI, REST APIs, React Native, Tailwind CSS")
    add_text(doc, theme, "Data and Infrastructure: PostgreSQL, MongoDB, Redis, SQLite, Prisma, Convex, Docker, Git, AWS, Vercel")
    add_text(doc, theme, "AI and Data: scikit-learn, TensorFlow, PyTorch, Pandas, NumPy, computer vision, model evaluation, feature engineering")

    add_section(doc, theme, "Professional Experience")
    add_role(doc, theme, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
    add_bullets(doc, theme, [
        "Contribute to application development and AI-enabled, agentic software systems.",
        "Work across implementation, integration, testing, and delivery using modern web and backend technologies.",
    ])
    add_role(doc, theme, "Software Engineer", "Hello World Communications Ltd", "Chattogram, Bangladesh", "Aug 2024-Present")
    add_bullets(doc, theme, [
        "Build, review, deploy, and maintain full-stack web applications for client and internal use across frontend, backend, and database layers.",
        "Collaborate with cross-functional stakeholders on requirements, implementation, production releases, and application support.",
        "Apply React, Next.js, Node.js, Python, PostgreSQL, and AI integrations to scalable application-development work.",
    ])
    add_role(doc, theme, "Freelance Developer", "Self-employed", "Remote", "Mar 2021-2025")
    add_bullets(doc, theme, [
        "Deliver web applications from requirements and interface design through backend implementation, deployment, and ongoing support.",
        "Built digital products for small businesses, nonprofit and academic organizations, and event teams using modern JavaScript and Python stacks.",
    ])
    page_break(doc)
    add_section(doc, theme, "Selected Projects")
    add_project_set(doc, theme, "software")

    add_section(doc, theme, "Education")
    add_role(doc, theme, "B.Sc. (Hons.) in Computer Science and Engineering", "BGC Trust University Bangladesh", "Chattogram, Bangladesh", "Expected Dec 2026")
    add_text(doc, theme, "Currently in the 8th semester. Focus: software development, machine learning, computer vision, and applied AI.")

    add_section(doc, theme, "Selected Publications")
    add_publication(doc, theme, "AgriScan: Next.js powered cross-platform solution for automated plant disease diagnosis and crop health management", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00169-7")
    add_publication(doc, theme, "Comparing pre-trained models for efficient leaf disease detection: a study on custom CNN", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00137-1")
    add_publication(doc, theme, "Next-Generation K-Means Clustering: Mojo-Driven Performance for Big Data", "International Journal of Intelligent Information Systems", "2025", "10.11648/j.ijiis.20251401.12")
    add_publication(doc, theme, "Architectures for AI-Driven Visual Assistance: Evaluating Server-Mediated Mobile and Direct Access Desktop Client Implementations", "Lecture Notes in Networks and Systems", "2026", "10.1007/978-3-032-15764-5_45")

    add_section(doc, theme, "Selected Certifications")
    add_text(doc, theme, "IBM Applied Data Science with Python; IBM Deep Learning with TensorFlow; Redis for Python Developers; Cisco Python Essentials 1 and 2; HackerRank Problem Solving (Intermediate)")
    add_leadership_and_honors(doc, theme)
    return doc


def research_resume(theme: Theme) -> Document:
    doc = configure_doc(theme)
    add_header(doc, theme, "Machine Learning Engineer | AI and Agentic Systems")
    add_section(doc, theme, "Research Profile")
    add_text(doc, theme, "Applied machine-learning researcher and software engineer working across computer vision, tabular modeling, healthcare analytics, agricultural AI, LLM inference, and high-performance computing. Author or co-author of ten source-backed research works, including two IEEE COMPAS papers, two SpringerOpen journal articles, and a 2026 Springer conference chapter. Builds reproducible Python pipelines and production-facing web systems that translate experiments into usable tools.")
    add_competition_highlights(doc, theme)

    add_section(doc, theme, "Research and Engineering Skills")
    add_text(doc, theme, "Machine Learning: scikit-learn, TensorFlow, PyTorch, XGBoost, LightGBM, CNNs, transfer learning, SVM, clustering, feature engineering, cross-validation, model evaluation")
    add_text(doc, theme, "Engineering: Python, TypeScript, Next.js, React, Node.js, FastAPI, Django, PostgreSQL, MongoDB, Redis, Docker, Git, Zig, Bun")
    add_text(doc, theme, "Research Practice: reproducible experiments, fold-safe validation, benchmark design, data pipelines, academic writing, literature review, publication preparation")

    add_section(doc, theme, "Engineering Experience")
    add_role(doc, theme, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
    add_bullets(doc, theme, [
        "Contribute to application development and AI-enabled, agentic software systems.",
        "Work across implementation, integration, testing, and delivery using modern web and backend technologies.",
    ])
    add_role(doc, theme, "Software Engineer", "Hello World Communications Ltd", "Chattogram, Bangladesh", "Aug 2024-Present")
    add_bullets(doc, theme, [
        "Develop and deploy scalable applications and AI-enabled features using TypeScript, React, Next.js, Node.js, Python, and PostgreSQL.",
        "Bridge experimental software and production application requirements through code review, system design, deployment, and support.",
    ])

    add_section(doc, theme, "Selected Research Projects")
    add_project_set(doc, theme, "research")

    page_break(doc)
    add_section(doc, theme, "Publications")
    publications = [
        ("Architectures for AI-Driven Visual Assistance: Evaluating Server-Mediated Mobile and Direct Access Desktop Client Implementations", "Lecture Notes in Networks and Systems", "2026", "10.1007/978-3-032-15764-5_45"),
        ("Enhancing Agricultural Diagnostics: Advanced Training of Pre-Trained CNN Models for Paddy Leaf Disease Detection", "Machine Learning Research", "2025", "10.11648/j.mlr.20251001.11"),
        ("Next-Generation K-Means Clustering: Mojo-Driven Performance for Big Data", "International Journal of Intelligent Information Systems", "2025", "10.11648/j.ijiis.20251401.12"),
        ("Efficient Malware Classification Using Multiprocessing and Bag-of-Words Vectorization", "Advances in Networks", "2025", "10.11648/j.net.20251201.12"),
        ("Fine-tuning LLaMA 2 interference: a comparative study of language implementations for optimal efficiency", "arXiv", "2025", "10.48550/arXiv.2502.01651"),
        ("AgriScan: Next.js powered cross-platform solution for automated plant disease diagnosis and crop health management", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00169-7"),
        ("Application of Machine Learning K-Means Clustering and Linear Regression in Determining the Risk Level of Pulmonary Tuberculosis", "IEEE COMPAS", "2024", "10.1109/COMPAS60761.2024.10796963"),
        ("Enhancing Cardiovascular Risk Prediction Using Support Vector Machines and Advanced Machine Learning Algorithms", "IEEE COMPAS", "2024", "10.1109/COMPAS60761.2024.10796805"),
        ("Comparing pre-trained models for efficient leaf disease detection: a study on custom CNN", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00137-1"),
        ("Comparative Performance Evaluation of Classical Machine Learning and Quantum SVM for Heart Disease Prediction using a Quantum-Featured Dataset", "Sonargaon University Journal", "2025", "https://su.edu.bd/web_assets/journal/journal_five/journal3.pdf"),
    ]
    for item in publications:
        add_publication(doc, theme, *item)

    add_section(doc, theme, "Education")
    add_role(doc, theme, "B.Sc. (Hons.) in Computer Science and Engineering", "BGC Trust University Bangladesh", "Chattogram, Bangladesh", "Expected Dec 2026")
    add_text(doc, theme, "Currently in the 8th semester. Research interests: applied machine learning, computer vision, agricultural AI, healthcare analytics, LLM systems, and agentic systems.")

    add_section(doc, theme, "Research Profiles")
    add_text(doc, theme, "ORCID: 0009-0007-7512-1893 | Google Scholar: scholar.google.com/citations?user=gRkTVYEAAAAJ | GitHub: github.com/Seyamalam")
    add_leadership_and_honors(doc, theme)
    return doc


def universal_resume(theme: Theme) -> Document:
    doc = configure_doc(theme)
    add_header(doc, theme, "Software Engineer | Web, App, Backend, ML/AI")
    add_section(doc, theme, "Professional Summary")
    add_text(doc, theme, "Software engineer, full-stack developer, and applied AI researcher who turns technical ideas into production-ready products and reproducible experiments. Professional experience since 2021 across TypeScript, Python, React, Next.js, Node.js, backend services, data systems, and machine learning. Builds client applications, open-source developer tools, AI agents, and research software; author or co-author of ten source-backed research works.")
    add_competition_highlights(doc, theme)

    add_section(doc, theme, "Core Skills")
    add_text(doc, theme, "Languages: TypeScript, JavaScript, Python, SQL, Go, Zig")
    add_text(doc, theme, "Product Engineering: React, Next.js, Node.js, Django, FastAPI, REST APIs, React Native, Tailwind CSS")
    add_text(doc, theme, "Data and Infrastructure: PostgreSQL, MongoDB, Redis, SQLite, Prisma, Convex, Docker, Git, AWS, Vercel")
    add_text(doc, theme, "AI and Research: scikit-learn, TensorFlow, PyTorch, XGBoost, LightGBM, computer vision, agent evaluation, reproducible experiments")

    add_section(doc, theme, "Professional Experience")
    add_role(doc, theme, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
    add_bullets(doc, theme, [
        "Contribute to application development and AI-enabled, agentic software systems.",
        "Work across implementation, integration, testing, and delivery using modern web and backend technologies.",
    ])
    add_role(doc, theme, "Software Engineer", "Hello World Communications Ltd", "Chattogram, Bangladesh", "Aug 2024-Present")
    add_bullets(doc, theme, [
        "Build, review, deploy, and maintain full-stack web applications for client and internal use across frontend, backend, database, and AI-integration layers.",
        "Collaborate with cross-functional stakeholders on requirements, implementation, production releases, and ongoing application support.",
    ])
    add_role(doc, theme, "Freelance Developer", "Self-employed", "Remote", "Mar 2021-2025")
    add_bullets(doc, theme, [
        "Delivered web applications from requirements and interface design through backend implementation, deployment, and ongoing support.",
        "Built digital products for small businesses, nonprofit and academic organizations, and event teams using modern JavaScript and Python stacks.",
    ])
    add_section(doc, theme, "Selected Projects")
    add_project_set(doc, theme, "universal")

    add_section(doc, theme, "Education")
    add_role(doc, theme, "B.Sc. (Hons.) in Computer Science and Engineering", "BGC Trust University Bangladesh", "Chattogram, Bangladesh", "Expected Dec 2026")
    add_text(doc, theme, "Currently in the 8th semester. Focus: software engineering, machine learning, computer vision, applied AI, and research-driven product development.")

    add_section(doc, theme, "Selected Research")
    add_publication(doc, theme, "Architectures for AI-Driven Visual Assistance: Evaluating Server-Mediated Mobile and Direct Access Desktop Client Implementations", "Lecture Notes in Networks and Systems", "2026", "10.1007/978-3-032-15764-5_45")
    add_publication(doc, theme, "AgriScan: Next.js powered cross-platform solution for automated plant disease diagnosis and crop health management", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00169-7")
    add_publication(doc, theme, "Next-Generation K-Means Clustering: Mojo-Driven Performance for Big Data", "International Journal of Intelligent Information Systems", "2025", "10.11648/j.ijiis.20251401.12")

    add_section(doc, theme, "Selected Certifications")
    add_text(doc, theme, "IBM Applied Data Science with Python; IBM Deep Learning with TensorFlow; Redis for Python Developers; Cisco Python Essentials 1 and 2; HackerRank Problem Solving (Intermediate)")
    add_leadership_and_honors(doc, theme)
    return doc


def compact_resume(theme: Theme) -> Document:
    doc = configure_doc(theme)
    add_header(doc, theme, "Web, App and Backend Developer | AI Systems")
    add_section(doc, theme, "Summary")
    add_text(doc, theme, "Software engineer and 8th-semester CSE student building web and mobile applications, backend services, Python/TypeScript developer tools, and applied ML systems. Currently at Agentic Institute and Hello World Communications Ltd.")
    add_competition_highlights(doc, theme, compact=True)
    add_section(doc, theme, "Core Skills")
    add_text(doc, theme, "TypeScript, JavaScript, Python, SQL, React, Next.js, Node.js, Django, FastAPI, Tailwind CSS, PostgreSQL, MongoDB, Redis, Prisma, Convex, Docker, Git, Vercel, TensorFlow, PyTorch, scikit-learn")
    add_section(doc, theme, "Experience")
    add_role(doc, theme, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
    add_bullets(doc, theme, ["Contribute to application development and AI-enabled, agentic software systems across implementation, integration, testing, and delivery."])
    add_role(doc, theme, "Software Engineer", "Hello World Communications Ltd", "Chattogram", "Aug 2024-Present")
    add_bullets(doc, theme, ["Build, review, deploy, and support full-stack and AI-enabled applications with cross-functional teams using React/Next.js, Node.js, Python, and PostgreSQL."])
    add_role(doc, theme, "Freelance Developer", "Self-employed", "Remote", "Mar 2021-2025")
    add_bullets(doc, theme, ["Deliver web applications for business, nonprofit, academic, and event use - from requirements and interface work through backend development and deployment."])
    add_section(doc, theme, "Selected Projects")
    add_project_set(doc, theme, "fullstack")
    add_section(doc, theme, "Education and Research")
    add_text(doc, theme, "B.Sc. (Hons.) Computer Science and Engineering, BGC Trust University Bangladesh - 8th semester; expected Dec 2026")
    add_text(doc, theme, "Ten source-backed research works, including IEEE COMPAS papers, SpringerOpen journal articles, and a 2026 Springer conference chapter. ORCID: 0009-0007-7512-1893")
    add_leadership_and_honors(doc, theme)
    return doc


def pdf_styles(theme: Theme):
    leading = theme.body_size * (1.12 if theme.compact else 1.18)
    section_fill = None
    section_color = theme.accent
    section_padding = 0
    if theme is MODERN:
        section_fill, section_padding = "E8F0FA", 4
    elif theme is COMPACT:
        section_fill, section_color, section_padding = theme.accent, "FFFFFF", 4
    elif theme is UNIVERSAL:
        section_fill, section_padding = "EAF0F5", 4
    return {
        "_theme": theme,
        "name": ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=theme.title_size,
                               textColor=HexColor("#" + theme.accent), leading=theme.title_size + 1,
                               alignment=TA_CENTER if theme is MODERN else TA_LEFT, spaceAfter=0),
        "target": ParagraphStyle("target", fontName="Helvetica-Bold", fontSize=9.5 if theme.compact else 10.5,
                                 textColor=HexColor("#" + theme.muted), leading=12,
                                 alignment=TA_CENTER if theme is MODERN else TA_LEFT, spaceAfter=2),
        "contact": ParagraphStyle("contact", fontName="Helvetica", fontSize=8.1 if theme.compact else 8.7,
                                  textColor=HexColor("#" + theme.muted), leading=10.2,
                                  alignment=TA_CENTER if theme is MODERN else TA_LEFT,
                                  spaceAfter=6 if theme.compact else 8),
        "h1": ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=10.7 if theme.compact else 11.6,
                             textColor=HexColor("#" + section_color), leading=14,
                             spaceBefore=5 if theme.compact else 7, spaceAfter=4,
                             keepWithNext=True, backColor=HexColor("#" + section_fill) if section_fill else None,
                             borderPadding=section_padding),
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=theme.body_size,
                               textColor=HexColor("#" + theme.ink), leading=leading,
                               spaceAfter=3 if theme.compact else 4),
        "role": ParagraphStyle("role", fontName="Helvetica-Bold", fontSize=theme.body_size + 0.2,
                               textColor=HexColor("#" + theme.ink), leading=leading,
                               spaceBefore=2, spaceAfter=1, keepWithNext=True),
        "bullet": ParagraphStyle("bullet", fontName="Helvetica", fontSize=theme.body_size,
                                 textColor=HexColor("#" + theme.ink), leading=leading,
                                 leftIndent=12 if theme.compact else 15, firstLineIndent=-7,
                                 bulletIndent=0, spaceAfter=1.5 if theme.compact else 2.5),
    }


def para(text: str, style, links: bool = False):
    escaped = text.replace("&", "&amp;")
    return Paragraph(escaped, style)


def pdf_header(story, styles, theme: Theme, target: str):
    story.append(Paragraph(NAME, styles["name"]))
    story.append(Paragraph(target.upper() if theme is MODERN else target, styles["target"]))
    story.append(Paragraph(CONTACT.replace("  |  ", " &nbsp;|&nbsp; "), styles["contact"]))
    story.append(HRFlowable(width="100%", thickness=2.2 if theme in (CLASSIC, UNIVERSAL) else 1.2,
                            color=HexColor("#" + theme.accent), spaceBefore=0, spaceAfter=4))


def pdf_section(story, styles, title):
    theme = styles["_theme"]
    if theme is CLASSIC:
        story.append(HRFlowable(width="100%", thickness=1.1, color=HexColor("#" + theme.accent),
                                spaceBefore=6, spaceAfter=2))
    story.append(Paragraph(title.upper(), styles["h1"]))


def pdf_text(story, styles, text):
    story.append(para(text, styles["body"]))


def pdf_role(story, styles, role, org, place, dates):
    story.append(Paragraph(f"{role} | {org} <font color='#66717D'>| {place} | {dates}</font>", styles["role"]))


def pdf_bullets(story, styles, items):
    for item in items:
        story.append(Paragraph("• " + item.replace("&", "&amp;"), styles["bullet"]))


def pdf_project(story, styles, name, stack, text, url=None):
    accent = styles["_theme"].accent
    link = f" | <link href='{url}' color='#{accent}'>{url.replace('https://', '')}</link>" if url else ""
    story.append(Paragraph(f"{name} <font color='#66717D'>| {stack}</font>{link}", styles["role"]))
    pdf_bullets(story, styles, [text])


def pdf_publication(story, styles, title, venue, year, doi):
    is_url = doi.startswith("http")
    href = doi if is_url else f"https://doi.org/{doi}"
    label = doi.replace("https://", "") if is_url else doi
    accent = styles["_theme"].accent
    story.append(Paragraph(f"• <i>“{title}.”</i> {venue}, {year}. {'Link' if is_url else 'DOI'}: <link href='{href}' color='#{accent}'>{label}</link>", styles["bullet"]))


def pdf_project_set(story, styles, key):
    for project in PROJECT_SETS[key]:
        pdf_project(story, styles, *project)


def pdf_competition_highlights(story, styles, *, compact=False):
    pdf_section(story, styles, "Selected Wins")
    pdf_bullets(story, styles, PRIMARY_WINS[:2] if compact else PRIMARY_WINS)


def pdf_leadership_and_honors(story, styles):
    pdf_section(story, styles, "Leadership and Additional Honors")
    pdf_text(story, styles, LEADERSHIP)
    pdf_bullets(story, styles, ADDITIONAL_HONORS)


def build_pdf(theme: Theme, kind: str, path: Path) -> None:
    styles = pdf_styles(theme)
    doc = SimpleDocTemplate(
        str(path), pagesize=letter,
        rightMargin=theme.margins * inch, leftMargin=theme.margins * inch,
        topMargin=theme.margins * inch, bottomMargin=theme.margins * inch,
        title=f"{NAME} - {kind}", author=NAME,
    )
    s = []
    if kind == "software":
        pdf_header(s, styles, theme, "Software Engineer | Web, App and Backend Systems")
        pdf_section(s, styles, "Professional Summary")
        pdf_text(s, styles, "Software engineer building production web applications, backend services, mobile systems, and AI-enabled products across TypeScript, Next.js, Node.js, Python, Go, PostgreSQL, and cloud infrastructure. Ships client platforms, resilient distributed prototypes, and open-source developer tooling from architecture through deployment.")
        pdf_competition_highlights(s, styles)
        pdf_section(s, styles, "Technical Skills")
        for t in ["<b>Languages:</b> TypeScript, JavaScript, Python, SQL, Go, Zig", "<b>Application Development:</b> React, Next.js, Node.js, Django, FastAPI, REST APIs, React Native, Tailwind CSS", "<b>Data and Infrastructure:</b> PostgreSQL, MongoDB, Redis, SQLite, Prisma, Convex, Docker, Git, AWS, Vercel", "<b>AI and Data:</b> scikit-learn, TensorFlow, PyTorch, Pandas, NumPy, computer vision, model evaluation, feature engineering"]: s.append(Paragraph(t, styles["body"]))
        pdf_section(s, styles, "Professional Experience")
        pdf_role(s, styles, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
        pdf_bullets(s, styles, ["Contribute to application development and AI-enabled, agentic software systems.", "Work across implementation, integration, testing, and delivery using modern web and backend technologies."])
        pdf_role(s, styles, "Software Engineer", "Hello World Communications Ltd", "Chattogram, Bangladesh", "Aug 2024-Present")
        pdf_bullets(s, styles, ["Build, review, deploy, and maintain full-stack web applications for client and internal use across frontend, backend, and database layers.", "Collaborate with cross-functional stakeholders on requirements, implementation, production releases, and application support.", "Apply React, Next.js, Node.js, Python, PostgreSQL, and AI integrations to scalable application-development work."])
        pdf_role(s, styles, "Freelance Developer", "Self-employed", "Remote", "Mar 2021-2025")
        pdf_bullets(s, styles, ["Deliver web applications from requirements and interface design through backend implementation, deployment, and ongoing support.", "Built digital products for small businesses, nonprofit and academic organizations, and event teams using modern JavaScript and Python stacks."])
        s.append(PageBreak())
        pdf_section(s, styles, "Selected Projects")
        pdf_project_set(s, styles, "software")
        pdf_section(s, styles, "Education")
        pdf_role(s, styles, "B.Sc. (Hons.) in Computer Science and Engineering", "BGC Trust University Bangladesh", "Chattogram, Bangladesh", "Expected Dec 2026")
        pdf_text(s, styles, "Currently in the 8th semester. Focus: software development, machine learning, computer vision, and applied AI.")
        pdf_section(s, styles, "Selected Publications")
        pdf_publication(s, styles, "AgriScan: Next.js powered cross-platform solution for automated plant disease diagnosis and crop health management", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00169-7")
        pdf_publication(s, styles, "Comparing pre-trained models for efficient leaf disease detection: a study on custom CNN", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00137-1")
        pdf_publication(s, styles, "Next-Generation K-Means Clustering: Mojo-Driven Performance for Big Data", "International Journal of Intelligent Information Systems", "2025", "10.11648/j.ijiis.20251401.12")
        pdf_publication(s, styles, "Architectures for AI-Driven Visual Assistance: Evaluating Server-Mediated Mobile and Direct Access Desktop Client Implementations", "Lecture Notes in Networks and Systems", "2026", "10.1007/978-3-032-15764-5_45")
        pdf_section(s, styles, "Selected Certifications")
        pdf_text(s, styles, "IBM Applied Data Science with Python; IBM Deep Learning with TensorFlow; Redis for Python Developers; Cisco Python Essentials 1 and 2; HackerRank Problem Solving (Intermediate)")
        pdf_leadership_and_honors(s, styles)
    elif kind == "research":
        pdf_header(s, styles, theme, "Machine Learning Engineer | AI and Agentic Systems")
        pdf_section(s, styles, "Research Profile")
        pdf_text(s, styles, "Applied machine-learning researcher and software engineer working across computer vision, tabular modeling, healthcare analytics, agricultural AI, LLM inference, and high-performance computing. Author or co-author of ten source-backed research works, including two IEEE COMPAS papers, two SpringerOpen journal articles, and a 2026 Springer conference chapter. Builds reproducible Python pipelines and production-facing web systems that translate experiments into usable tools.")
        pdf_competition_highlights(s, styles)
        pdf_section(s, styles, "Research and Engineering Skills")
        for t in ["<b>Machine Learning:</b> scikit-learn, TensorFlow, PyTorch, XGBoost, LightGBM, CNNs, transfer learning, SVM, clustering, feature engineering, cross-validation, model evaluation", "<b>Engineering:</b> Python, TypeScript, Next.js, React, Node.js, FastAPI, Django, PostgreSQL, MongoDB, Redis, Docker, Git, Zig, Bun", "<b>Research Practice:</b> reproducible experiments, fold-safe validation, benchmark design, data pipelines, academic writing, literature review, publication preparation"]: s.append(Paragraph(t, styles["body"]))
        pdf_section(s, styles, "Engineering Experience")
        pdf_role(s, styles, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
        pdf_bullets(s, styles, ["Contribute to application development and AI-enabled, agentic software systems.", "Work across implementation, integration, testing, and delivery using modern web and backend technologies."])
        pdf_role(s, styles, "Software Engineer", "Hello World Communications Ltd", "Chattogram, Bangladesh", "Aug 2024-Present")
        pdf_bullets(s, styles, ["Develop and deploy scalable applications and AI-enabled features using TypeScript, React, Next.js, Node.js, Python, and PostgreSQL.", "Bridge experimental software and production application requirements through code review, system design, deployment, and support."])
        pdf_section(s, styles, "Selected Research Projects")
        pdf_project_set(s, styles, "research")
        s.append(PageBreak())
        pdf_section(s, styles, "Publications")
        pubs = [("Architectures for AI-Driven Visual Assistance: Evaluating Server-Mediated Mobile and Direct Access Desktop Client Implementations", "Lecture Notes in Networks and Systems", "2026", "10.1007/978-3-032-15764-5_45"), ("Enhancing Agricultural Diagnostics: Advanced Training of Pre-Trained CNN Models for Paddy Leaf Disease Detection", "Machine Learning Research", "2025", "10.11648/j.mlr.20251001.11"), ("Next-Generation K-Means Clustering: Mojo-Driven Performance for Big Data", "International Journal of Intelligent Information Systems", "2025", "10.11648/j.ijiis.20251401.12"), ("Efficient Malware Classification Using Multiprocessing and Bag-of-Words Vectorization", "Advances in Networks", "2025", "10.11648/j.net.20251201.12"), ("Fine-tuning LLaMA 2 interference: a comparative study of language implementations for optimal efficiency", "arXiv", "2025", "10.48550/arXiv.2502.01651"), ("AgriScan: Next.js powered cross-platform solution for automated plant disease diagnosis and crop health management", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00169-7"), ("Application of Machine Learning K-Means Clustering and Linear Regression in Determining the Risk Level of Pulmonary Tuberculosis", "IEEE COMPAS", "2024", "10.1109/COMPAS60761.2024.10796963"), ("Enhancing Cardiovascular Risk Prediction Using Support Vector Machines and Advanced Machine Learning Algorithms", "IEEE COMPAS", "2024", "10.1109/COMPAS60761.2024.10796805"), ("Comparing pre-trained models for efficient leaf disease detection: a study on custom CNN", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00137-1"), ("Comparative Performance Evaluation of Classical Machine Learning and Quantum SVM for Heart Disease Prediction using a Quantum-Featured Dataset", "Sonargaon University Journal", "2025", "https://su.edu.bd/web_assets/journal/journal_five/journal3.pdf")]
        for item in pubs: pdf_publication(s, styles, *item)
        pdf_section(s, styles, "Education")
        pdf_role(s, styles, "B.Sc. (Hons.) in Computer Science and Engineering", "BGC Trust University Bangladesh", "Chattogram, Bangladesh", "Expected Dec 2026")
        pdf_text(s, styles, "Currently in the 8th semester. Research interests: applied machine learning, computer vision, agricultural AI, healthcare analytics, LLM systems, and agentic systems.")
        pdf_section(s, styles, "Research Profiles")
        pdf_text(s, styles, "ORCID: 0009-0007-7512-1893 | Google Scholar: scholar.google.com/citations?user=gRkTVYEAAAAJ | GitHub: github.com/Seyamalam")
        pdf_leadership_and_honors(s, styles)
    elif kind == "universal":
        pdf_header(s, styles, theme, "Software Engineer | Web, App, Backend, ML/AI")
        pdf_section(s, styles, "Professional Summary")
        pdf_text(s, styles, "Software engineer, full-stack developer, and applied AI researcher who turns technical ideas into production-ready products and reproducible experiments. Professional experience since 2021 across TypeScript, Python, React, Next.js, Node.js, backend services, data systems, and machine learning. Builds client applications, open-source developer tools, AI agents, and research software; author or co-author of ten source-backed research works.")
        pdf_competition_highlights(s, styles)
        pdf_section(s, styles, "Core Skills")
        for t in ["<b>Languages:</b> TypeScript, JavaScript, Python, SQL, Go, Zig", "<b>Product Engineering:</b> React, Next.js, Node.js, Django, FastAPI, REST APIs, React Native, Tailwind CSS", "<b>Data and Infrastructure:</b> PostgreSQL, MongoDB, Redis, SQLite, Prisma, Convex, Docker, Git, AWS, Vercel", "<b>AI and Research:</b> scikit-learn, TensorFlow, PyTorch, XGBoost, LightGBM, computer vision, agent evaluation, reproducible experiments"]:
            s.append(Paragraph(t, styles["body"]))
        pdf_section(s, styles, "Professional Experience")
        pdf_role(s, styles, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
        pdf_bullets(s, styles, ["Contribute to application development and AI-enabled, agentic software systems.", "Work across implementation, integration, testing, and delivery using modern web and backend technologies."])
        pdf_role(s, styles, "Software Engineer", "Hello World Communications Ltd", "Chattogram, Bangladesh", "Aug 2024-Present")
        pdf_bullets(s, styles, ["Build, review, deploy, and maintain full-stack web applications for client and internal use across frontend, backend, database, and AI-integration layers.", "Collaborate with cross-functional stakeholders on requirements, implementation, production releases, and ongoing application support."])
        pdf_role(s, styles, "Freelance Developer", "Self-employed", "Remote", "Mar 2021-2025")
        pdf_bullets(s, styles, ["Delivered web applications from requirements and interface design through backend implementation, deployment, and ongoing support.", "Built digital products for small businesses, nonprofit and academic organizations, and event teams using modern JavaScript and Python stacks."])
        pdf_section(s, styles, "Selected Projects")
        pdf_project_set(s, styles, "universal")
        pdf_section(s, styles, "Education")
        pdf_role(s, styles, "B.Sc. (Hons.) in Computer Science and Engineering", "BGC Trust University Bangladesh", "Chattogram, Bangladesh", "Expected Dec 2026")
        pdf_text(s, styles, "Currently in the 8th semester. Focus: software engineering, machine learning, computer vision, applied AI, and research-driven product development.")
        pdf_section(s, styles, "Selected Research")
        pdf_publication(s, styles, "Architectures for AI-Driven Visual Assistance: Evaluating Server-Mediated Mobile and Direct Access Desktop Client Implementations", "Lecture Notes in Networks and Systems", "2026", "10.1007/978-3-032-15764-5_45")
        pdf_publication(s, styles, "AgriScan: Next.js powered cross-platform solution for automated plant disease diagnosis and crop health management", "Journal of Electrical Systems and Information Technology", "2024", "10.1186/s43067-024-00169-7")
        pdf_publication(s, styles, "Next-Generation K-Means Clustering: Mojo-Driven Performance for Big Data", "International Journal of Intelligent Information Systems", "2025", "10.11648/j.ijiis.20251401.12")
        pdf_section(s, styles, "Selected Certifications")
        pdf_text(s, styles, "IBM Applied Data Science with Python; IBM Deep Learning with TensorFlow; Redis for Python Developers; Cisco Python Essentials 1 and 2; HackerRank Problem Solving (Intermediate)")
        pdf_leadership_and_honors(s, styles)
    else:
        pdf_header(s, styles, theme, "Web, App and Backend Developer | AI Systems")
        pdf_section(s, styles, "Summary")
        pdf_text(s, styles, "Software engineer and 8th-semester CSE student building web and mobile applications, backend services, Python/TypeScript developer tools, and applied ML systems. Currently at Agentic Institute and Hello World Communications Ltd.")
        pdf_competition_highlights(s, styles, compact=True)
        pdf_section(s, styles, "Core Skills")
        pdf_text(s, styles, "TypeScript, JavaScript, Python, SQL, React, Next.js, Node.js, Django, FastAPI, Tailwind CSS, PostgreSQL, MongoDB, Redis, Prisma, Convex, Docker, Git, Vercel, TensorFlow, PyTorch, scikit-learn")
        pdf_section(s, styles, "Experience")
        pdf_role(s, styles, "Software Engineer", "Agentic Institute", "Remote", "Dec 2025-Present")
        pdf_bullets(s, styles, ["Contribute to application development and AI-enabled, agentic software systems across implementation, integration, testing, and delivery."])
        pdf_role(s, styles, "Software Engineer", "Hello World Communications Ltd", "Chattogram", "Aug 2024-Present")
        pdf_bullets(s, styles, ["Build, review, deploy, and support full-stack and AI-enabled applications with cross-functional teams using React/Next.js, Node.js, Python, and PostgreSQL."])
        pdf_role(s, styles, "Freelance Developer", "Self-employed", "Remote", "Mar 2021-2025")
        pdf_bullets(s, styles, ["Deliver web applications for business, nonprofit, academic, and event use - from requirements and interface work through backend development and deployment."])
        pdf_section(s, styles, "Selected Projects")
        pdf_project_set(s, styles, "fullstack")
        pdf_section(s, styles, "Education and Research")
        pdf_text(s, styles, "B.Sc. (Hons.) Computer Science and Engineering, BGC Trust University Bangladesh - 8th semester; expected Dec 2026")
        pdf_text(s, styles, "Ten source-backed research works, including IEEE COMPAS papers, SpringerOpen journal articles, and a 2026 Springer conference chapter. ORCID: 0009-0007-7512-1893")
        pdf_leadership_and_honors(s, styles)
    doc.build(s)


def main() -> None:
    variants = [
        (CLASSIC, "software", software_resume),
        (MODERN, "research", research_resume),
        (COMPACT, "compact", compact_resume),
        (UNIVERSAL, "universal", universal_resume),
    ]
    artifacts = []
    for theme, kind, builder in variants:
        base = f"Touhidul_Alam_Seyam_{theme.key}"
        document = builder(theme)
        document.core_properties.title = f"{NAME} - {theme.key.replace('_', ' ')}"
        document.core_properties.subject = "ATS-compatible resume"
        document.core_properties.author = NAME
        document.core_properties.keywords = "software engineer, web development, app development, backend, machine learning, AI, agentic systems, TypeScript, Python"
        docx_path = DIST / f"{base}.docx"
        pdf_path = DIST / f"{base}.pdf"
        document.save(docx_path)
        build_pdf(theme, kind, pdf_path)
        artifacts.extend((docx_path, pdf_path))
        print(f"created {base}.docx and {base}.pdf ({theme.preset})")

    pack_path = DIST / "Touhidul_Alam_Seyam_ATS_Resume_Pack.zip"
    with ZipFile(pack_path, "w", compression=ZIP_DEFLATED) as archive:
        for artifact in artifacts:
            archive.write(artifact, arcname=artifact.name)
    print(f"created {pack_path.name} with {len(artifacts)} files")


if __name__ == "__main__":
    main()
